"""Unit checks for the multimodal document feature:
  • _decide_file_route — the caption + smart-default disambiguation (import vs
    analyze vs note). This is the core of "how do I tell import from analysis".
  • document_service — kind detection, image magic-byte sniffing, content-block
    building, the scanned-PDF→vision fallback, and (banking compliance) that
    extracted text is PII-redacted BEFORE it becomes a Claude content block.

Run:  venv/bin/python tests/document_routing_check.py
"""
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Repoint DB to a throwaway path before importing (matches the rest of the suite;
# nothing here touches the DB, but keeps the real data file safe regardless).
import config
config.DATABASE_PATH = "/tmp/yordamchi_docroute_test.db"

import document_service  # noqa: E402
import handlers  # noqa: E402

PASS = 0
FAIL = 0
FAILED = []


def check(name, cond, detail=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  ✅ {name}")
    else:
        FAIL += 1
        FAILED.append(name)
        print(f"  ❌ {name}   {detail}")


def test_route_matrix():
    print("\n[ _decide_file_route — caption + smart default ]")
    r = handlers._decide_file_route
    # No caption → type-based default.
    check("no caption, excel → import", r("", "excel") == "import")
    check("no caption, csv → import", r("", "csv") == "import")
    check("no caption, pdf → analyze", r("", "pdf") == "analyze")
    check("no caption, docx → analyze", r("", "docx") == "analyze")
    check("no caption, image → analyze", r("", "image") == "analyze")
    # Caption says import (and kind is importable) → import.
    check("'import qil' + pdf → import", r("import qil", "pdf") == "import")
    check("'ro'yxatni yukla' + excel → import", r("ro'yxatni yukla", "excel") == "import")
    # Import hint but a non-importable kind → analyze (can't import an image).
    check("'import' + image → analyze", r("import qil", "image") == "analyze")
    # Analysis-style captions → analyze (even on an importable kind).
    check("'qisqacha' + pdf → analyze", r("qisqacha aytib ber", "pdf") == "analyze")
    check("'tahlil qil' + excel → analyze", r("tahlil qil", "excel") == "analyze")
    # 'vazifaga qo'sh' is NOT an import hint — Claude extracts via actions instead.
    check("'vazifaga qo'sh' + pdf → analyze", r("muhim sanalarni vazifaga qo'sh", "pdf") == "analyze")
    # Note hints → note (regardless of kind).
    check("'qayd qil' + pdf → note", r("qayd qil", "pdf") == "note")
    check("'inbox' + image → note", r("inbox ga", "image") == "note")


def test_doc_cache():
    print("\n[ _doc_remember — file cache stores kind+mime, bounded ]")
    handlers._DOC_FLIP_CACHE.clear()
    handlers._doc_remember("uid1", "fileid1", "contract.pdf", "pdf", "application/pdf")
    e = handlers._DOC_FLIP_CACHE.get("uid1")
    check("entry stored", e is not None)
    check("file_id stored", bool(e) and e.get("file_id") == "fileid1")
    check("kind stored", bool(e) and e.get("kind") == "pdf")
    check("mime stored", bool(e) and e.get("mime") == "application/pdf")
    for i in range(handlers._DOC_FLIP_CACHE_MAX + 5):
        handlers._doc_remember(f"u{i}", f"f{i}", "x.pdf", "pdf")
    check("cache bounded at max", len(handlers._DOC_FLIP_CACHE) <= handlers._DOC_FLIP_CACHE_MAX)


def test_detect_kind():
    print("\n[ document_service.detect_kind ]")
    d = document_service.detect_kind
    check("contract.pdf → pdf", d("contract.pdf", "") == "pdf")
    check("mime application/pdf → pdf", d("x", "application/pdf") == "pdf")
    check("memo.docx → docx", d("memo.docx", "") == "docx")
    check("tasks.xlsx → excel", d("tasks.xlsx", "") == "excel")
    check("data.csv → csv", d("data.csv", "") == "csv")
    check("scan.jpg → image", d("scan.JPG", "") == "image")
    check("mime image/png → image", d("blob", "image/png") == "image")
    check("notes.txt → text", d("notes.txt", "") == "text")
    check("archive.zip → other", d("archive.zip", "application/zip") == "other")


def test_image_media_type():
    print("\n[ document_service._image_media_type (magic bytes) ]")
    m = document_service._image_media_type
    check("JPEG magic", m(b"\xff\xd8\xff\xe0rest", None) == "image/jpeg")
    check("PNG magic", m(b"\x89PNG\r\n\x1a\nrest", None) == "image/png")
    check("GIF magic", m(b"GIF89a....", None) == "image/gif")
    check("WEBP magic", m(b"RIFF\x00\x00\x00\x00WEBPrest", None) == "image/webp")
    check("mime hint wins", m(b"garbage", "image/png") == "image/png")
    check("unsupported (BMP) → None", m(b"BM\x00\x00", "image/bmp") is None)


def test_build_blocks_image():
    print("\n[ build_content_blocks — image (vision) ]")
    blocks, meta = document_service.build_content_blocks(
        b"\x89PNG\r\n\x1a\n" + b"\x00" * 50, "image", "scan.png", "image/png")
    check("one block", len(blocks) == 1)
    check("type=image", blocks[0]["type"] == "image")
    check("media_type png", blocks[0]["source"]["media_type"] == "image/png")
    check("base64 source", blocks[0]["source"]["type"] == "base64")
    check("mode=vision", meta["mode"] == "vision")


def test_build_blocks_pdf_scanned_fallback():
    print("\n[ build_content_blocks — non-text PDF falls back to vision document block ]")
    # Not a real PDF → pypdf fails → text empty → scanned/vision document block.
    blocks, meta = document_service.build_content_blocks(
        b"%PDF-1.4 garbage not really a pdf", "pdf", "scan.pdf", "application/pdf")
    check("one block", len(blocks) == 1)
    check("type=document", blocks[0]["type"] == "document")
    check("media_type pdf", blocks[0]["source"]["media_type"] == "application/pdf")
    check("mode=vision", meta["mode"] == "vision")


def test_build_blocks_docx_redacts():
    print("\n[ build_content_blocks — docx extraction + PII redaction (compliance) ]")
    try:
        from docx import Document
    except ImportError:
        check("python-docx available", False, "python-docx not installed")
        return
    doc = Document()
    doc.add_paragraph("Shartnoma bo'yicha to'lov.")
    doc.add_paragraph("Aloqa: +998901234567 raqamiga qo'ng'iroq qiling.")
    buf = io.BytesIO()
    doc.save(buf)
    blocks, meta = document_service.build_content_blocks(
        buf.getvalue(), "docx", "memo.docx", "")
    text = blocks[0]["text"] if blocks and blocks[0]["type"] == "text" else ""
    check("mode=text", meta.get("mode") == "text")
    check("content extracted", "Shartnoma" in text)
    check("phone redacted (raw gone)", "998901234567" not in text)
    check("redaction marker present", "[PHONE-REDACTED]" in text)
    check("redacted count > 0", meta.get("redacted", 0) >= 1)


def test_document_errors():
    print("\n[ build_content_blocks — guards raise DocumentError ]")
    DocumentError = document_service.DocumentError
    try:
        document_service.build_content_blocks(b"", "image", "x.png", "image/png")
        check("empty file raises", False)
    except DocumentError:
        check("empty file raises", True)
    try:
        oversize = b"x" * (document_service.MAX_FILE_BYTES + 1)
        document_service.build_content_blocks(oversize, "text", "big.txt", "text/plain")
        check("oversize raises", False)
    except DocumentError:
        check("oversize raises", True)
    try:
        document_service.build_content_blocks(b"data", "other", "a.zip", "application/zip")
        check("unsupported kind raises", False)
    except DocumentError:
        check("unsupported kind raises", True)


def main():
    print("=" * 60)
    print("DOCUMENT / MULTIMODAL ROUTING CHECKS")
    print("=" * 60)
    test_route_matrix()
    test_doc_cache()
    test_detect_kind()
    test_image_media_type()
    test_build_blocks_image()
    test_build_blocks_pdf_scanned_fallback()
    test_build_blocks_docx_redacts()
    test_document_errors()
    print("\n" + "=" * 60)
    print(f"RESULT: {PASS} passed, {FAIL} failed")
    if FAILED:
        print("FAILED: " + ", ".join(FAILED))
    print("=" * 60)
    sys.exit(1 if FAIL else 0)


if __name__ == "__main__":
    main()
