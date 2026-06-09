"""protocol_doc.py — Agrobank bayonnoma Word generatori testi.

Docx generatsiya qilib, python-docx bilan qayta o'qiydi va dizaynni tekshiradi
(ko'k shapka/jadval boshi, zebra, struktura, Lotin↔Kiril). Faqat python-docx kerak.
"""

import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import protocol_doc  # noqa: E402

_PASS = 0
_FAIL = 0
_FAILED: list[str] = []


def check(name: str, ok: bool, extra: str = "") -> None:
    global _PASS, _FAIL
    print(f"  [{'✓' if ok else '✗'}] {name}" + ("" if ok else f"  {extra}"))
    if ok:
        _PASS += 1
    else:
        _FAIL += 1
        _FAILED.append(name)


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _cell_fill(cell):
    from docx.oxml.ns import qn
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def main() -> None:
    from docx import Document

    meeting = {
        "title": "Vazirlik bilan hamkorlik bo'yicha ish uchrashuvi",
        "datetime_start": "2026-06-10T14:00:00+05:00",
        "participants": ["H. Begmatjonov", "M. Rustamov", "Vazirlik vakillari"],
        "agenda": "Tadbirga tayyorgarlik: videokontent, lokatsiyalar, homiylik.",
    }
    tasks = [
        {"assignee": "Muhammadzoir", "title": "Videorolik tayyorlash", "deadline": "2026-06-11T09:00:00+05:00"},
        {"assignee": "", "title": "Markaziy sektor bilan muzokara", "deadline": None},
        {"assignee": "Farrux", "title": "Qo'shimcha materiallar olish", "deadline": None},
    ]

    print("── build_fields ──")
    f = protocol_doc.build_fields(meeting, "", tasks, settings=None)
    check("org default = AGROBANK", '"AGROBANK" ATB' == f["org_name"], f["org_name"])
    check("author default = Maqsud Rustamov", f["author"] == "Maqsud Rustamov")
    check("date = 2026-yil 10-iyun", f["date_str"] == "2026-yil 10-iyun", f["date_str"])
    check("3 participants", len(f["participants"]) == 3)
    check("missing assignee → '—'", f["tasks"][1]["assignee"] == "—")
    # settings override
    f2 = protocol_doc.build_fields(meeting, "", tasks, settings={"org_dept": "Xazina boshqarmasi"})
    check("settings override dept", f2["org_dept"] == "Xazina boshqarmasi")

    print("\n── build_docx (LATIN) ──")
    blob = protocol_doc.build_docx(f, script="latin")
    check("docx bytes returned (non-trivial)", isinstance(blob, bytes) and len(blob) > 3000, str(len(blob)))
    doc = Document(io.BytesIO(blob))
    txt = _all_text(doc)
    check("header: AGROBANK", "AGROBANK" in txt)
    check("title: BAYONNOMA № ____", "BAYONNOMA № ____" in txt)
    check("subtitle: (ish uchrashuvi)", "(ish uchrashuvi)" in txt)
    check("place: Toshkent shahri", "Toshkent shahri" in txt)
    check("date in body", "2026-yil 10-iyun" in txt)
    check("Mavzu present", "Mavzu:" in txt)
    check("Ishtirokchilar heading", "Uchrashuv ishtirokchilari:" in txt)
    check("Kun tartibi heading", "Kun tartibi:" in txt)
    check("task header label", "Topshiriq mazmuni" in txt and "Mas'ul shaxs(lar)" in txt)
    check("signatures heading", "Ishtirokchilar imzosi:" in txt)
    check("compiler line", "Bayonnomani tuzdi:" in txt)
    check("missing deadline → '—' in table", "—" in txt)

    # tables: place/date(2) + task table(4 cols) + signature(2)
    check("3 tables (place/date, tasks, signatures)", len(doc.tables) == 3, str(len(doc.tables)))
    task_tbl = doc.tables[1]
    check("task table: header + 3 rows", len(task_tbl.rows) == 4, str(len(task_tbl.rows)))
    check("task header navy fill", _cell_fill(task_tbl.rows[0].cells[0]) == protocol_doc.NAVY,
          str(_cell_fill(task_tbl.rows[0].cells[0])))
    # zebra: row idx 2 (even task #2) shaded, row idx 1 (#1) not
    check("zebra: row #2 shaded #F2F5FB", _cell_fill(task_tbl.rows[2].cells[0]) == protocol_doc.ZEBRA,
          str(_cell_fill(task_tbl.rows[2].cells[0])))
    check("zebra: row #1 NOT shaded", _cell_fill(task_tbl.rows[1].cells[0]) in (None, "auto"),
          str(_cell_fill(task_tbl.rows[1].cells[0])))
    # column order: title in col 2, assignee in col 3
    check("col order: title→col2, assignee→col3",
          "Videorolik" in task_tbl.rows[1].cells[1].text and "Muhammadzoir" in task_tbl.rows[1].cells[2].text)

    print("\n── build_docx (CYRILLIC) ──")
    blob_c = protocol_doc.build_docx(f, script="cyrillic")
    doc_c = Document(io.BytesIO(blob_c))
    txt_c = _all_text(doc_c)
    check("cyrillic: БАЁННОМА", "БАЁННОМА" in txt_c, txt_c[:60])
    check("cyrillic: Мавзу", "Мавзу" in txt_c)
    check("cyrillic: Тошкент шаҳри", "Тошкент шаҳри" in txt_c)
    check("cyrillic: no latin 'Mavzu'", "Mavzu:" not in txt_c)
    # design preserved under cyrillic
    check("cyrillic: 3 tables", len(doc_c.tables) == 3)
    check("cyrillic: navy header fill intact",
          _cell_fill(doc_c.tables[1].rows[0].cells[0]) == protocol_doc.NAVY)

    print("\n── build_pdf ──")
    if not protocol_doc.pdf_available():
        check("reportlab o'rnatilgan (PDF)", False, "reportlab yo'q — PDF testlari o'tkazib yuborildi")
    else:
        pdf = protocol_doc.build_pdf(f, script="latin")
        check("PDF magic %PDF", isinstance(pdf, bytes) and pdf[:4] == b"%PDF", str(pdf[:8]))
        check("PDF non-trivial size", len(pdf) > 2000, str(len(pdf)))
        pdf_c = protocol_doc.build_pdf(f, script="cyrillic")
        check("PDF (cyrillic) %PDF", pdf_c[:4] == b"%PDF")
        try:
            from pypdf import PdfReader
            rl = PdfReader(io.BytesIO(pdf))
            ptxt = "\n".join((pg.extract_text() or "") for pg in rl.pages)
            check("PDF text: AGROBANK", "AGROBANK" in ptxt, ptxt[:80])
            check("PDF text: BAYONNOMA", "BAYONNOMA" in ptxt)
            check("PDF text: Toshkent", "Toshkent" in ptxt)
            rc = PdfReader(io.BytesIO(pdf_c))
            ctxt = "\n".join((pg.extract_text() or "") for pg in rc.pages)
            check("PDF (cyrillic) text: Кирилcha gliflar", ("Тошкент" in ctxt or "БАЁННОМА" in ctxt or "Мавзу" in ctxt), ctxt[:80])
        except ImportError:
            check("pypdf mavjud (matn tekshiruvi)", False, "pypdf yo'q — matn tekshiruvi o'tkazib yuborildi")

    print("\n" + "=" * 50)
    print(f"NATIJA:  ✅ {_PASS} o'tdi   ❌ {_FAIL} yiqildi")
    if _FAILED:
        print("Yiqilganlar: " + ", ".join(_FAILED))
    print("=" * 50)
    sys.exit(1 if _FAIL else 0)


if __name__ == "__main__":
    main()
