"""Rasmiy bayonnoma hujjat generatori — Agrobank korporativ shabloni.

Word (.docx) va PDF chiqishini bitta strukturalangan `fields` dict'idan chizadi.
Lotin/Kiril — `script` parametri orqali (translit.py).

Dizayn (foydalanuvchi shabloni "Bayonnoma shabloni.docx"):
  - Korporativ ko'k #1F3864 (sarlavhalar, jadval boshi, № ustun, imzo yorliqlari)
  - Och-ko'k zebra #F2F5FB (topshiriqlar jadvali navbatma-navbat qatorlari)
  - Kulrang #595959 / #808080 (boshqarma subtitle, izoh, "imzo")
  - Arial, US Letter (chap 2.54 / o'ng 1.3 sm)
  - Tuzilma: shapka → joy/sana → Mavzu → Ishtirokchilar → Kun tartibi →
    topshiriqlar jadvali → imzolar → "Bayonnomani tuzdi"

handlers.py minimal o'zgartirilsin (race-ni kamaytirish) — bu modul mustaqil:
faqat `translit`ga bog'liq; topshiriqlar ro'yxati tashqarida ajratiladi va
`tasks` orqali uzatiladi (har biri {assignee, title, deadline}).
"""

from __future__ import annotations

import io

import translit

# ── Palette ──
NAVY = "1F3864"
ZEBRA = "F2F5FB"
GRAY = "595959"
GRAY2 = "808080"
WHITE = "FFFFFF"

_UZ_MONTHS = ["yanvar", "fevral", "mart", "aprel", "may", "iyun",
              "iyul", "avgust", "sentabr", "oktabr", "noyabr", "dekabr"]

# Sozlamalar (database.get_settings()) bilan almashtiriladigan standart qiymatlar.
DEFAULTS = {
    "org_name": '"AGROBANK" ATB',
    "org_dept": "Marketing boshqarmasi",
    "org_place": "Toshkent shahri",
    "protocol_author": "Maqsud Rustamov",
    "protocol_subtitle": "(ish uchrashuvi)",
}


def _fmt_deadline(iso) -> str:
    """Jadval 'Muddat' ustuni: 'DD-oy' yoki '—' (noma'lum)."""
    if not iso:
        return "—"
    try:
        from datetime import datetime
        dt = datetime.fromisoformat(str(iso))
        return f"{dt.day}-{_UZ_MONTHS[dt.month - 1]}"
    except (ValueError, TypeError, IndexError):
        return str(iso)


def build_fields(meeting: dict | None, protocol_text: str, tasks, settings: dict | None = None) -> dict:
    """Meeting + settings + (oldindan ajratilgan) tasks → strukturalangan maydonlar.

    `tasks` — har biri {assignee, title, deadline} bo'lgan ro'yxat (handlers tomonidan
    `_proto_tasks_from_actions` / `_proto_tasks_from_text` orqali ajratiladi)."""
    from datetime import datetime
    s = settings or {}
    m = meeting or {}

    def cfg(key: str) -> str:
        v = s.get(key)
        if isinstance(v, str) and v.strip():
            return v.strip()
        return DEFAULTS.get(key, "")

    date_str = ""
    ds = m.get("datetime_start")
    if ds:
        try:
            dt = datetime.fromisoformat(str(ds))
            date_str = f"{dt.year}-yil {dt.day}-{_UZ_MONTHS[dt.month - 1]}"
        except (ValueError, TypeError, IndexError):
            date_str = ""

    participants = [str(p).strip() for p in (m.get("participants") or []) if str(p).strip()]

    norm_tasks = []
    for t in (tasks or []):
        norm_tasks.append({
            "assignee": (t.get("assignee") or "").strip() or "—",
            "title": (t.get("title") or "").strip() or "—",
            "deadline": t.get("deadline"),
        })

    return {
        "org_name": cfg("org_name"),
        "org_dept": cfg("org_dept"),
        "subtitle": cfg("protocol_subtitle"),
        "place": cfg("org_place"),
        "date_str": date_str,
        "mavzu": (m.get("title") or "").strip(),
        "participants": participants,
        "agenda": (m.get("agenda") or "").strip(),
        "tasks": norm_tasks,
        "author": cfg("protocol_author"),
    }


_HEADERS = ["№", "Topshiriq mazmuni", "Mas'ul shaxs(lar)", "Muddat"]
_IZOH = ('Izoh: "—" bilan belgilangan kataklardagi mas\'ul shaxslar va muddatlar '
         "uchrashuv yakunida aniqlashtirilib, to'ldiriladi.")


def build_docx(fields: dict, script: str = "latin") -> bytes:
    """Agrobank korporativ bayonnomasini Word (.docx) bayt sifatida qaytaradi."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def tr(text: str) -> str:
        return translit.transliterate(text or "", script)

    def rgb(h: str) -> RGBColor:
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    c_navy, c_gray, c_gray2, c_white = rgb(NAVY), rgb(GRAY), rgb(GRAY2), rgb(WHITE)

    def shade(cell, fill: str) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def no_borders(tbl) -> None:
        tblPr = tbl._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
            borders.append(e)
        tblPr.append(borders)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.59)
    sec.page_height = Cm(27.94)  # US Letter
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(1.3)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    _rf = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        _rf.set(qn(a), "Arial")

    C, J, R = WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.RIGHT

    def para(align=None, before=4, after=4, indent=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        if indent is not None:
            p.paragraph_format.left_indent = Cm(indent)
        return p

    def run(p, text, size=11, bold=False, italic=False, color=None):
        r = p.add_run(tr(text))
        r.bold = bold
        r.italic = italic
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), "Arial")
        if color is not None:
            r.font.color.rgb = color
        return r

    # ── Shapka ──
    run(para(C, before=0, after=0), fields.get("org_name") or "", size=13, bold=True, color=c_navy)
    if fields.get("org_dept"):
        run(para(C, before=0, after=0), fields["org_dept"], size=10, color=c_gray)
    run(para(C, before=8, after=0), "BAYONNOMA № ____", size=15, bold=True)
    if fields.get("subtitle"):
        run(para(C, before=0, after=6), fields["subtitle"], size=10, italic=True, color=c_gray)

    # ── Joy / sana (chegrasiz 2 ustun) ──
    pd = doc.add_table(rows=1, cols=2)
    no_borders(pd)
    lc, rc = pd.rows[0].cells
    lc.width = Cm(8.25)
    rc.width = Cm(8.25)
    lc.paragraphs[0].text = ""
    run(lc.paragraphs[0], fields.get("place") or "")
    rp = rc.paragraphs[0]
    rp.text = ""
    rp.alignment = R
    run(rp, fields.get("date_str") or "")

    # ── Mavzu ──
    if fields.get("mavzu"):
        p = para(J, before=8)
        run(p, "Mavzu: ", bold=True)
        run(p, fields["mavzu"])

    # ── Ishtirokchilar ──
    if fields.get("participants"):
        run(para(J, before=8, after=2), "Uchrashuv ishtirokchilari:", bold=True, color=c_navy)
        for nm in fields["participants"]:
            run(para(J, before=0, after=0, indent=0.5), f"—  {nm}")

    # ── Kun tartibi ──
    if fields.get("agenda"):
        run(para(J, before=8, after=2), "Kun tartibi:", bold=True, color=c_navy)
        run(para(J), fields["agenda"])

    # ── Topshiriqlar ──
    run(para(J, before=8, after=2), "Muhokama qilindi va quyidagi topshiriqlar belgilandi:",
        bold=True, color=c_navy)
    run(para(J, before=0, after=4), _IZOH, size=9, italic=True, color=c_gray2)

    tasks = fields.get("tasks") or []
    tbl = doc.add_table(rows=1, cols=4)
    try:
        tbl.style = "Table Grid"
    except KeyError:
        pass
    widths = [Cm(1.2), Cm(9.55), Cm(3.25), Cm(3.75)]
    hdr = tbl.rows[0].cells
    for i, h in enumerate(_HEADERS):
        shade(hdr[i], NAVY)
        hp = hdr[i].paragraphs[0]
        hp.text = ""
        hp.alignment = C
        run(hp, h, size=10, bold=True, color=c_white)
    for idx, t in enumerate(tasks, 1):
        cells = tbl.add_row().cells
        if idx % 2 == 0:
            for cc in cells:
                shade(cc, ZEBRA)
        vals = [str(idx), t.get("title") or "—", t.get("assignee") or "—", _fmt_deadline(t.get("deadline"))]
        for i, v in enumerate(vals):
            cp = cells[i].paragraphs[0]
            cp.text = ""
            if i == 0:
                cp.alignment = C
                run(cp, v, size=10, bold=True, color=c_navy)
            else:
                run(cp, v, size=10)
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    # ── Imzolar ──
    if fields.get("participants"):
        run(para(J, before=10, after=2), "Ishtirokchilar imzosi:", bold=True, color=c_navy)
        sg = doc.add_table(rows=0, cols=2)
        no_borders(sg)
        for nm in fields["participants"]:
            cells = sg.add_row().cells
            cells[0].width = Cm(9.88)
            cells[1].width = Cm(6.63)
            cells[0].paragraphs[0].text = ""
            run(cells[0].paragraphs[0], nm)
            sp = cells[1].paragraphs[0]
            sp.text = ""
            sp.alignment = R
            run(sp, "______________  imzo", size=9, color=c_gray2)

    # ── Bayonnomani tuzdi ──
    p = para(J, before=12)
    run(p, "Bayonnomani tuzdi: ", bold=True)
    run(p, (fields.get("author") or "") + "  _______________________")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF (reportlab — lazy; Debian/macOS Unicode TTF; graceful degradation) ──
_PDF_FONTS = None  # cache: (regular_name, bold_name)

_TTF_REGULAR = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",                 # Debian/Ubuntu
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/System/Library/Fonts/Supplemental/Arial.ttf",                    # macOS
    "/Library/Fonts/Arial.ttf",
]
_TTF_BOLD = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
    "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
    "/Library/Fonts/Arial Bold.ttf",
]


def pdf_available() -> bool:
    """True if reportlab is importable (PDF can be generated)."""
    try:
        import reportlab  # noqa: F401
        return True
    except ImportError:
        return False


def _register_pdf_fonts():
    """Register a Cyrillic-capable Unicode TTF (so Kiril renders, not boxes).
    Falls back to Helvetica (Latin-only) when no TTF is found — Cyrillic PDFs
    then need a system font installed. Cached (re-registration errors)."""
    global _PDF_FONTS
    if _PDF_FONTS is not None:
        return _PDF_FONTS
    import os
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    reg = next((p for p in _TTF_REGULAR if os.path.exists(p)), None)
    if reg:
        bold = next((p for p in _TTF_BOLD if os.path.exists(p)), None)
        try:
            pdfmetrics.registerFont(TTFont("PDoc", reg))
            pdfmetrics.registerFont(TTFont("PDoc-Bold", bold or reg))
            _PDF_FONTS = ("PDoc", "PDoc-Bold")
            return _PDF_FONTS
        except Exception:
            pass
    _PDF_FONTS = ("Helvetica", "Helvetica-Bold")  # Latin-only fallback
    return _PDF_FONTS


def build_pdf(fields: dict, script: str = "latin") -> bytes:
    """Agrobank korporativ bayonnomasini PDF bayt sifatida qaytaradi (Word bilan
    bir xil dizayn). reportlab yo'q bo'lsa aniq xato beradi (handler ushlaydi)."""
    try:
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT, TA_LEFT
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Table,
                                        TableStyle, Spacer)
    except ImportError as e:
        raise RuntimeError(
            "PDF uchun 'reportlab' o'rnatilmagan. Serverda bir marta:\n"
            "  ~/yordamchi_ai/venv/bin/python -m pip install reportlab"
        ) from e

    from xml.sax.saxutils import escape as _esc

    FONT, FONT_B = _register_pdf_fonts()
    navy = colors.HexColor("#" + NAVY)
    zebra = colors.HexColor("#" + ZEBRA)
    gray = colors.HexColor("#" + GRAY)
    gray2 = colors.HexColor("#" + GRAY2)
    white = colors.white
    grid = colors.HexColor("#BBBBBB")

    def tx(text: str) -> str:
        return _esc(translit.transliterate(text or "", script))

    def st(name, size, bold=False, color=colors.black, align=TA_LEFT, lead=None):
        return ParagraphStyle(name, fontName=(FONT_B if bold else FONT), fontSize=size,
                              textColor=color, alignment=align, leading=lead or size * 1.3)

    s_org = st("org", 13, bold=True, color=navy, align=TA_CENTER)
    s_dept = st("dept", 10, color=gray, align=TA_CENTER)
    s_title = st("title", 15, bold=True, align=TA_CENTER)
    s_sub = st("sub", 10, color=gray, align=TA_CENTER)
    s_place = st("place", 11)
    s_date = st("date", 11, align=TA_RIGHT)
    s_head = st("head", 11, bold=True, color=navy)
    s_body = st("body", 11, align=TA_JUSTIFY)
    s_izoh = st("izoh", 9, color=gray2)
    s_th = st("th", 10, bold=True, color=white, align=TA_CENTER)
    s_num = st("num", 10, bold=True, color=navy, align=TA_CENTER)
    s_cell = st("cell", 10)
    s_sig = st("sig", 9, color=gray2, align=TA_RIGHT)

    def P(text, style):
        return Paragraph(tx(text), style)

    story = []
    story.append(P(fields.get("org_name") or "", s_org))
    if fields.get("org_dept"):
        story.append(P(fields["org_dept"], s_dept))
    story.append(Spacer(1, 8))
    story.append(P("BAYONNOMA № ____", s_title))
    if fields.get("subtitle"):
        story.append(P(fields["subtitle"], s_sub))
    story.append(Spacer(1, 8))

    pd = Table([[P(fields.get("place") or "", s_place), P(fields.get("date_str") or "", s_date)]],
               colWidths=[8.25 * cm, 8.25 * cm])
    pd.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                            ("TOPPADDING", (0, 0), (-1, -1), 0),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 2)]))
    story.append(pd)
    story.append(Spacer(1, 6))

    if fields.get("mavzu"):
        story.append(Paragraph(f'<b>{_esc(translit.transliterate("Mavzu: ", script))}</b>'
                               + tx(fields["mavzu"]), s_body))
    if fields.get("participants"):
        story.append(Spacer(1, 4))
        story.append(P("Uchrashuv ishtirokchilari:", s_head))
        for nm in fields["participants"]:
            story.append(P(f"—  {nm}", s_body))
    if fields.get("agenda"):
        story.append(Spacer(1, 4))
        story.append(P("Kun tartibi:", s_head))
        story.append(P(fields["agenda"], s_body))

    story.append(Spacer(1, 6))
    story.append(P("Muhokama qilindi va quyidagi topshiriqlar belgilandi:", s_head))
    story.append(P(_IZOH, s_izoh))
    story.append(Spacer(1, 4))

    data = [[P(h, s_th) for h in _HEADERS]]
    for idx, t in enumerate(fields.get("tasks") or [], 1):
        data.append([
            P(str(idx), s_num),
            P(t.get("title") or "—", s_cell),
            P(t.get("assignee") or "—", s_cell),
            P(_fmt_deadline(t.get("deadline")), s_cell),
        ])
    tbl = Table(data, colWidths=[1.2 * cm, 9.55 * cm, 3.25 * cm, 3.75 * cm], repeatRows=1)
    ts = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), navy),
        ("GRID", (0, 0), (-1, -1), 0.5, grid),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
    ])
    for i in range(2, len(data), 2):  # even task rows → zebra
        ts.add("BACKGROUND", (0, i), (-1, i), zebra)
    tbl.setStyle(ts)
    story.append(tbl)

    if fields.get("participants"):
        story.append(Spacer(1, 10))
        story.append(P("Ishtirokchilar imzosi:", s_head))
        sig_rows = [[P(nm, s_cell), P("______________  imzo", s_sig)]
                    for nm in fields["participants"]]
        sg = Table(sig_rows, colWidths=[9.88 * cm, 6.63 * cm])
        sg.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0),
                                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                                ("TOPPADDING", (0, 0), (-1, -1), 3),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
        story.append(sg)

    story.append(Spacer(1, 12))
    story.append(Paragraph(f'<b>{_esc(translit.transliterate("Bayonnomani tuzdi: ", script))}</b>'
                           + tx((fields.get("author") or "") + "  _______________________"), s_body))

    buf = io.BytesIO()
    SimpleDocTemplate(
        buf, pagesize=(21.59 * cm, 27.94 * cm),
        leftMargin=2.54 * cm, rightMargin=1.3 * cm,
        topMargin=2.54 * cm, bottomMargin=2.54 * cm,
        title="Bayonnoma",
    ).build(story)
    return buf.getvalue()
