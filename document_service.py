"""Document & image ingestion — turns an uploaded file into Anthropic content
blocks for multimodal analysis.

Compliance posture (banking): text-extractable documents (PDF/DOCX/XLSX/CSV/TXT)
are extracted locally and **PII-redacted** (redaction.py) BEFORE anything leaves
the process — same trust boundary as typed text. Images and scanned PDFs are sent
as native vision blocks; their pixels can't be pre-scrubbed, so callers should
surface meta["mode"] == "vision" when that matters.

Pure module: no Telegram / Claude coupling. handlers downloads the bytes and
hands them here; claude_service consumes the returned content blocks.
"""

import base64
import io
import logging
from typing import Optional

import redaction

logger = logging.getLogger(__name__)

# Telegram Bot API getFile caps downloads at 20 MB — stay under it. Also well
# under Claude's 32 MB PDF ceiling.
MAX_FILE_BYTES = 18 * 1024 * 1024
# Claude PDF document-block page ceiling (only matters for the scanned/vision path).
MAX_PDF_PAGES = 100
# A PDF that yields fewer than this many extracted characters is treated as
# scanned (image-only) and routed to the vision path for OCR.
_MIN_PDF_TEXT_CHARS = 200
# Cap extracted text so a 100-page contract can't blow the context window / cost.
# ~60k chars ≈ 15k tokens — generous for a summary while staying bounded.
_MAX_TEXT_CHARS = 60_000
# Image media types Claude accepts.
_SUPPORTED_IMAGE_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp"}


class DocumentError(Exception):
    """Carries a human-readable (O'zbek) reason a file can't be processed.
    handlers surfaces .args[0] directly to the principal."""


def detect_kind(file_name: Optional[str], mime: Optional[str]) -> str:
    """Classify an upload into a processing kind from its name and/or MIME type.
    Returns one of: pdf, docx, excel, csv, image, text, other."""
    name = (file_name or "").lower()
    mime = (mime or "").lower()
    if name.endswith(".pdf") or mime == "application/pdf":
        return "pdf"
    if name.endswith(".docx") or "wordprocessingml" in mime:
        return "docx"
    if name.endswith((".xlsx", ".xlsm")) or "spreadsheetml" in mime:
        return "excel"
    if name.endswith(".csv") or mime == "text/csv":
        return "csv"
    if name.endswith((".jpg", ".jpeg", ".png", ".gif", ".webp")) or mime.startswith("image/"):
        return "image"
    if name.endswith((".txt", ".md", ".log")) or mime.startswith("text/"):
        return "text"
    return "other"


def _image_media_type(data: bytes, mime: Optional[str]) -> Optional[str]:
    """Resolve a Claude-supported image media type from the MIME hint or magic
    bytes. Returns None for unsupported formats (BMP/TIFF/HEIC/...)."""
    mime = (mime or "").lower()
    if mime in _SUPPORTED_IMAGE_TYPES:
        return mime
    if data[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if data[:6] in (b"GIF87a", b"GIF89a"):
        return "image/gif"
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    return None


def extract_pdf_text(data: bytes) -> tuple[str, int]:
    """Return (concatenated_text, page_count). Empty text for scanned PDFs."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = len(reader.pages)
    text = "\n".join((p.extract_text() or "") for p in reader.pages)
    return text, pages


def extract_docx_text(data: bytes) -> str:
    """Plain text from a .docx — paragraphs plus table cells (kept in reading
    order well enough for summarisation)."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_excel_text(data: bytes) -> str:
    """Dump every sheet's non-empty cells to a readable text grid."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"# {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _text_block(text: str) -> tuple[list[dict], dict]:
    """Truncate, redact, and wrap extracted text as a single text content block."""
    truncated = len(text) > _MAX_TEXT_CHARS
    if truncated:
        text = text[:_MAX_TEXT_CHARS]
    redacted, n = redaction.redact(text)
    return (
        [{"type": "text", "text": redacted}],
        {"mode": "text", "redacted": n, "chars": len(text), "truncated": truncated},
    )


def build_content_blocks(
    data: bytes, kind: str, file_name: str = "", mime: str = ""
) -> tuple[list[dict], dict]:
    """Turn raw file bytes into Anthropic content blocks for a user message.

    Returns (blocks, meta). meta keys: mode ('text'|'vision'), redacted (int),
    chars (int), pages (int, PDFs only), truncated (bool).

    Raises DocumentError with an O'zbek reason when the file can't be turned into
    blocks (too big, empty, unsupported format, unreadable)."""
    if not data:
        raise DocumentError("Fayl bo'sh.")
    if len(data) > MAX_FILE_BYTES:
        raise DocumentError(
            f"Fayl juda katta ({len(data) // (1024 * 1024)} MB). "
            f"{MAX_FILE_BYTES // (1024 * 1024)} MB gacha qabul qila olaman."
        )

    if kind == "image":
        media = _image_media_type(data, mime)
        if not media:
            raise DocumentError(
                "Bu rasm formatini o'qiy olmayman. JPEG, PNG, GIF yoki WEBP yuboring."
            )
        b64 = base64.standard_b64encode(data).decode("ascii")
        block = {"type": "image", "source": {"type": "base64", "media_type": media, "data": b64}}
        return [block], {"mode": "vision", "redacted": 0, "chars": 0, "pages": 1, "truncated": False}

    if kind == "pdf":
        try:
            text, pages = extract_pdf_text(data)
        except Exception as e:
            logger.warning("pypdf extract failed (%s) — falling back to vision", e)
            text, pages = "", 0
        if len(text.strip()) >= _MIN_PDF_TEXT_CHARS:
            blocks, meta = _text_block(text)
            meta["pages"] = pages
            return blocks, meta
        # Scanned / image-only PDF → native document block (Claude OCRs it).
        if pages and pages > MAX_PDF_PAGES:
            raise DocumentError(
                f"Skan qilingan PDF juda uzun ({pages} bet). "
                f"{MAX_PDF_PAGES} betgacha tahlil qila olaman."
            )
        b64 = base64.standard_b64encode(data).decode("ascii")
        block = {
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
        }
        return [block], {"mode": "vision", "redacted": 0, "chars": 0,
                         "pages": pages, "truncated": False}

    if kind == "docx":
        try:
            text = extract_docx_text(data)
        except Exception as e:
            logger.warning("docx extract failed: %s", e)
            raise DocumentError("Word faylini o'qib bo'lmadi.")
        if not text.strip():
            raise DocumentError("Word faylida matn topilmadi.")
        return _text_block(text)

    if kind == "excel":
        try:
            text = extract_excel_text(data)
        except Exception as e:
            logger.warning("excel extract failed: %s", e)
            raise DocumentError("Excel faylini o'qib bo'lmadi.")
        if not text.strip():
            raise DocumentError("Excel faylida ma'lumot topilmadi.")
        return _text_block(text)

    if kind in ("text", "csv"):
        try:
            text = data.decode("utf-8-sig", errors="replace")
        except Exception:
            raise DocumentError("Matn faylini o'qib bo'lmadi.")
        if not text.strip():
            raise DocumentError("Fayl bo'sh.")
        return _text_block(text)

    raise DocumentError(
        "Bu fayl turini tahlil qila olmayman. PDF, Word, rasm yoki matn yuboring."
    )
